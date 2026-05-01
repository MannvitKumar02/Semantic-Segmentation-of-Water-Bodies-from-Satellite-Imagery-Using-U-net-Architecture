from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Water Body Segmentation — Pipeline Reference', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Dataset: Kaggle — franciscoescobar/satellite-images-of-water-bodies')
doc.add_paragraph('2841 image-mask pairs, folders: Images/ and Masks/, filenames match.')
doc.add_paragraph('Task: Binary segmentation — water (1) vs non-water (0)')

# Section 1
doc.add_heading('1. Setup', level=1)
doc.add_paragraph('Platform: Google Colab + GPU (T4)')
doc.add_paragraph('Libraries: PyTorch, torchvision, matplotlib, PIL, numpy')
doc.add_paragraph('Mount Google Drive to save model/results')

# Section 2
doc.add_heading('2. Dataset & Preprocessing', level=1)
items2 = [
    'Custom Dataset class (PyTorch)',
    'Resize image + mask to 256x256 (BILINEAR for image, NEAREST for mask)',
    'Normalize image with ImageNet mean/std: [0.485,0.456,0.406] / [0.229,0.224,0.225]',
    'Binarize mask: pixel > 127 → 1.0, else → 0.0, shape [1, H, W]',
    'Augmentation (train only, separate dataset instance): hflip, vflip, rotate 0/90/180/270, brightness/contrast jitter',
    'Split: 70% train / 15% val / 15% test (seed=42)',
    'DataLoaders: batch=8 for train/val, batch=1 for test',
]
for item in items2:
    doc.add_paragraph(item, style='List Bullet')

# Section 3
doc.add_heading('3. U-Net Architecture (built from scratch)', level=1)
items3 = [
    'DoubleConv: Conv2d → BN → ReLU → Conv2d → BN → ReLU (bias=False)',
    'EncoderBlock: DoubleConv → save skip → MaxPool2d(2)',
    'DecoderBlock: ConvTranspose2d(stride=2) → pad if needed → concat skip → DoubleConv',
    'UNet: 4 encoders [64,128,256,512] → bottleneck [1024] → 4 decoders → 1x1 Conv output',
    'Output: single channel logits (no sigmoid, applied at inference)',
]
for item in items3:
    doc.add_paragraph(item, style='List Bullet')

# Section 4
doc.add_heading('4. Loss Function', level=1)
items4 = [
    'DiceLoss: 1 - (2*intersection + smooth) / (pred + target + smooth), smooth=1.0',
    'BCEDiceLoss: 0.5 * BCEWithLogitsLoss + 0.5 * DiceLoss',
]
for item in items4:
    doc.add_paragraph(item, style='List Bullet')

# Section 5
doc.add_heading('5. Metrics', level=1)
items5 = [
    'iou_score(logits, targets, threshold=0.5): per-batch mean IoU with 1e-6 epsilon',
    'dice_score(logits, targets, threshold=0.5): per-batch mean Dice with 1e-6 epsilon',
]
for item in items5:
    doc.add_paragraph(item, style='List Bullet')

# Section 6
doc.add_heading('6. Training', level=1)
items6 = [
    'Optimizer: Adam, lr=1e-4',
    'Scheduler: ReduceLROnPlateau(mode=min, patience=5, factor=0.5)',
    'Epochs: 20',
    'Track: train_loss, val_loss, val_iou, val_dice per epoch',
    'Save best model to Drive when val_loss improves',
]
for item in items6:
    doc.add_paragraph(item, style='List Bullet')

# Section 7
doc.add_heading('7. Evaluation', level=1)
items7 = [
    'Load best model weights',
    'Run on test_loader (batch=1)',
    'Report: Test IoU and Test Dice (averaged over all test samples)',
]
for item in items7:
    doc.add_paragraph(item, style='List Bullet')

# Section 8
doc.add_heading('8. Visualization', level=1)
items8 = [
    'Plot 1: Training curves — train loss, val loss, val IoU (3 subplots)',
    'Plot 2: 6x3 grid — Input Image | Ground Truth | Prediction, with IoU per row',
    'Save results plot to Drive as segmentation_results.png',
]
for item in items8:
    doc.add_paragraph(item, style='List Bullet')

doc.save('basic_pipeline.docx')
print("basic_pipeline.docx created successfully!")
